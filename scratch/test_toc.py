import subprocess
import os
import sys
import docx
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    md_file = os.path.abspath('File Utili/Guida Web App.md')
    docx_file = os.path.abspath('File Utili/Guida Web App.docx')

    # Step 1: Run Pandoc with --toc --toc-depth=3
    cmd = ['pandoc', md_file, '--toc', '--toc-depth=3', '-o', docx_file]
    print("Running pandoc with --toc...", cmd)
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        print("Pandoc error:", res.stderr)
        sys.exit(1)

    doc = docx.Document(docx_file)

    # 1. Clean up manual w:br page breaks
    for p in list(doc.paragraphs):
        xml_str = p._p.xml
        if 'w:br' in xml_str and 'w:type="page"' in xml_str:
            for run in p.runs:
                for br in run._r.xpath('.//w:br[@w:type="page"]'):
                    br.getparent().remove(br)

    # 2. Format Cover Page (Title, Subtitle, Date, Logo)
    for p in doc.paragraphs:
        if p.style.name in ['Title', 'Subtitle', 'Date']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        contains_image = any(r._r.xpath('.//w:drawing') for r in p.runs)
        if contains_image:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(24)
            for run in p.runs:
                for inline in run._r.xpath('.//wp:inline'):
                    extent = inline.find('{http://schemas.openxmlformats.org/drawingml/2006/main}extent')
                    if extent is not None:
                        extent.set('cx', '2560320')
                        extent.set('cy', '731520')

    # 3. Locate the SDT (TOC) element and insert a "Sommario" Heading right before it
    body = doc._body._element
    sdts = body.xpath('.//w:sdt')
    if sdts:
        sdt = sdts[0]
        # Create a Heading 1 paragraph for "Sommario"
        h_p = parse_xml(r'''
            <w:p %s>
                <w:pPr>
                    <w:pStyle w:val="Heading1"/>
                    <w:pageBreakBefore/>
                    <w:spacing w:before="240" w:after="240"/>
                </w:pPr>
                <w:r>
                    <w:t>Sommario</w:t>
                </w:r>
            </w:p>
        ''' % nsdecls('w'))
        # Insert "Sommario" heading paragraph right before the sdt block
        sdt.getparent().insert(sdt.getparent().index(sdt), h_p)
        print("Successfully inserted 'Sommario' heading before TOC.")

    # 4. Set page_break_before = True on all Chapter Heading 1s
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            p.paragraph_format.page_break_before = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)

    doc.save(docx_file)
    print("Saved DOCX with Sommario heading and TOC field.")

if __name__ == '__main__':
    main()
