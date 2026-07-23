import subprocess
import os
import sys
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    md_file = os.path.abspath('File Utili/Guida Web App.md')
    docx_file = os.path.abspath('File Utili/Guida Web App.docx')

    # Step 1: Run Pandoc with --toc and --toc-depth=3 to generate Table of Contents (Sommario)
    cmd = ['pandoc', md_file, '--toc', '--toc-depth=3', '-o', docx_file]
    print("Running pandoc with --toc...", cmd)
    res = subprocess.run(cmd, capture_output=True, text=True)
    if res.returncode != 0:
        print("Pandoc error:", res.stderr)
        sys.exit(1)
    print("Pandoc output generated successfully.")

    # Step 2: Open generated DOCX
    doc = docx.Document(docx_file)
    body = doc._body._element

    # Clean up any manual page break tags inserted in paragraphs
    for p in list(doc.paragraphs):
        xml_str = p._p.xml
        if 'w:br' in xml_str and 'w:type="page"' in xml_str:
            for run in p.runs:
                for br in run._r.xpath('.//w:br[@w:type="page"]'):
                    br.getparent().remove(br)

    # 1. Find the paragraph containing the Logo image
    img_p = None
    for p in doc.paragraphs:
        if any(r._r.xpath('.//w:drawing') for r in p.runs):
            img_p = p
            break

    # 2. Find the Date paragraph (or Subtitle paragraph)
    date_p = None
    for p in doc.paragraphs:
        if p.style.name == 'Date':
            date_p = p
            break
        elif p.style.name == 'Subtitle' and not date_p:
            date_p = p

    # 3. Move the Logo image paragraph so it comes IMMEDIATELY AFTER Date paragraph on Page 1
    if img_p is not None and date_p is not None and img_p._p != date_p._p:
        # Move img_p._p element in XML to be right after date_p._p
        date_elem = date_p._p
        img_elem = img_p._p
        parent = date_elem.getparent()
        # Remove img_elem from current location
        parent.remove(img_elem)
        # Insert img_elem right after date_elem
        date_index = parent.index(date_elem)
        parent.insert(date_index + 1, img_elem)
        print("Successfully moved Logo image paragraph right below Date on Page 1.")

    # 4. Format Title block (Title, Subtitle, Date) & Logo image
    for p in doc.paragraphs:
        if p.style.name in ['Title', 'Subtitle', 'Date']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
        
        contains_image = any(r._r.xpath('.//w:drawing') for r in p.runs)
        if contains_image:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(24)
            p.paragraph_format.page_break_before = False
            
            # Set logo size to ~2.8 inches
            for run in p.runs:
                for inline in run._r.xpath('.//wp:inline'):
                    extent = inline.find('{http://schemas.openxmlformats.org/drawingml/2006/main}extent')
                    if extent is not None:
                        extent.set('cx', '2560320')
                        extent.set('cy', '731520')

    # 5. Locate the SDT (TOC) element and insert "Sommario" Heading 1 right before it
    sdts = body.xpath('.//w:sdt')
    if sdts:
        sdt = sdts[0]
        h_p = parse_xml(r'''
            <w:p %s>
                <w:pPr>
                    <w:pStyle w:val="Heading1"/>
                    <w:pageBreakBefore/>
                    <w:spacing w:before="240" w:after="240"/>
                </w:pPr>
                <w:r>
                    <w:t>Sommario</w:t>
                </w:r>
            </w:p>
        ''' % nsdecls('w'))
        sdt.getparent().insert(sdt.getparent().index(sdt), h_p)
        print("Inserted 'Sommario' heading before TOC field.")

    # 6. Set page_break_before = True for all Chapter Heading 1s
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1':
            p.paragraph_format.page_break_before = True
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)

    # Save document
    doc.save(docx_file)
    print("Successfully saved DOCX with Logo on Cover Page 1 under Title/Subtitle/Date.")

if __name__ == '__main__':
    main()
