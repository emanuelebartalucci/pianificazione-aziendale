import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re

def test_fix():
    docx_file = os.path.abspath('File Utili/Guida Web App.docx')
    doc = docx.Document(docx_file)

    # 1. Center Title, Subtitle, Date and Logo
    for p in doc.paragraphs:
        # Title block centering
        if p.style.name in ['Title', 'Subtitle', 'Date']:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Logo centering
        for r in p.runs:
            if r._r.xpath('.//w:drawing'):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(24)
                
                # Check logo width / size
                for inline in r._r.xpath('.//wp:inline'):
                    extent = inline.find('{http://schemas.openxmlformats.org/drawingml/2006/main}extent')
                    if extent is not None:
                        # Set width to ~2.5 inches (2286000 EMUs)
                        extent.set('cx', '2286000')
                        # Maintain aspect ratio (approx height 600000 EMUs)
                        extent.set('cy', '650000')

    # 2. Check font assignment for emoji characters
    # If a run contains emoji characters, set run.font.name = 'Segoe UI Emoji'
    emoji_pattern = re.compile(
        r'[\U0001F600-\U0001F64F'  # emoticons
        r'\U0001F300-\U0001F5FF'  # symbols & pictographs
        r'\U0001F680-\U0001F6FF'  # transport & map symbols
        r'\U0001F1E0-\U0001F1FF'  # flags
        r'\U0002700-\U00027BF'    # dingbats
        r'\U0001F900-\U0001F9FF'  # supplemental symbols
        r'\U0001FA70-\U0001FAFF'  # symbols and pictographs extended
        r'\u2100-\u26FF'          # misc symbols
        r'\u2700-\u27BF'
        r']+', flags=re.UNICODE
    )

    for p in doc.paragraphs:
        for r in p.runs:
            if emoji_pattern.search(r.text):
                r.font.name = 'Segoe UI Emoji'

    doc.save(docx_file)
    print("Saved test fix to docx.")

if __name__ == '__main__':
    test_fix()
