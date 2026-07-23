import docx

def verify():
    doc = docx.Document('File Utili/Guida Web App.docx')
    page_breaks = []
    empty_paragraphs = []

    for idx, p in enumerate(doc.paragraphs):
        xml = p._p.xml
        has_page_break = 'w:br' in xml and 'w:type="page"' in xml
        if has_page_break:
            page_breaks.append((idx, p.text[:30]))
        if not p.text.strip() and not any(r._r.xpath('.//w:drawing') for r in p.runs):
            empty_paragraphs.append(idx)

    print("=== VERIFICATION RESULTS ===")
    print("Explicit w:br page breaks:", page_breaks)
    print("Empty text paragraphs:", len(empty_paragraphs))
    print("Total paragraphs:", len(doc.paragraphs))

if __name__ == '__main__':
    verify()
